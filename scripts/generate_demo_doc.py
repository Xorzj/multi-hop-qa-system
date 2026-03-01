"""生成端到端 Demo 测试文档。

构造一个小型光网络技术文档，实体和关系清晰，
专门用于验证 知识图谱构建 → 多跳推理问答 全流程。

预期图谱:
    OTN ──派生──> SDH
    OTN ──包含──> DWDM
    SDH ──包含──> STM-16
    SDH ──包含──> STM-64
    SDH ──协同──> DWDM
    DWDM ──包含──> OADM
    DWDM ──依赖──> EDFA
    DWDM ──依赖──> 光纤
    EDFA ──属于──> 光放大器
    光纤 ──属于──> 传输介质

多跳路径:
    Q1: "OTN网络使用了什么光信号放大技术？"
        → OTN→包含→DWDM→依赖→EDFA→属于→光放大器 (3跳)
    Q2: "SDH体系与哪些光网络设备协同工作？"
        → SDH→协同→DWDM→包含→OADM (2跳)
    Q3: "OTN体系支持哪些传输速率等级？"
        → OTN→派生→SDH→包含→STM-16/STM-64 (2跳)
"""

from docx import Document
from docx.shared import Pt


def create_demo_document(output_path: str = "test_documents/demo.docx") -> None:
    """创建 Demo 测试文档。"""
    doc = Document()

    # ── 标题 ──
    doc.add_heading("光传输网络核心技术概述", level=0)

    # ── 第1章 OTN ──
    doc.add_heading("第一章 OTN光传送网", level=1)

    doc.add_heading("1.1 OTN概述", level=2)
    doc.add_paragraph(
        "OTN（光传送网）是新一代光传输体系架构，"
        "由国际电信联盟ITU-T定义。"
        "OTN在设计上继承并派生自SDH（同步数字体系）的管理理念，"
        "同时将DWDM（密集波分复用）技术纳入其体系框架之中。"
        "因此，OTN既包含了DWDM的波分复用能力，"
        "又保留了SDH的完善运维管理机制，"
        "是光网络领域的核心标准体系。"
    )

    doc.add_heading("1.2 OTN与SDH的关系", level=2)
    doc.add_paragraph(
        "OTN体系派生自SDH技术。"
        "SDH为OTN提供了完善的开销管理、保护倒换和性能监测能力。"
        "在实际网络部署中，OTN承载层通常与SDH客户层协同运行，"
        "实现端到端的电信级传输保障。"
    )

    # ── 第2章 SDH ──
    doc.add_heading("第二章 SDH同步数字体系", level=1)

    doc.add_heading("2.1 SDH概述", level=2)
    doc.add_paragraph(
        "SDH（同步数字体系）是一种时分复用传输标准，"
        "广泛应用于电信骨干网。"
        "SDH体系包含多个速率等级，"
        "其中STM-16提供2.5Gbps传输速率，"
        "STM-64提供10Gbps传输速率。"
        "这两种速率等级是SDH网络中最常用的配置。"
    )

    doc.add_heading("2.2 SDH与DWDM的协同", level=2)
    doc.add_paragraph(
        "在现代光网络中，SDH与DWDM技术协同工作。"
        "SDH负责业务的复用和调度，"
        "而DWDM负责在光纤上实现大容量波分传输。"
        "两者的协同使得单根光纤可以同时承载数十个SDH通道，"
        "大幅提升了网络传输容量。"
    )

    # ── 第3章 DWDM ──
    doc.add_heading("第三章 DWDM密集波分复用", level=1)

    doc.add_heading("3.1 DWDM概述", level=2)
    doc.add_paragraph(
        "DWDM（密集波分复用）是光纤通信的关键技术，"
        "通过在单根光纤中同时传输多个不同波长的光信号，"
        "实现传输容量的倍增。"
        "DWDM系统的正常运行依赖于光纤作为传输介质，"
        "同时依赖EDFA（掺铒光纤放大器）进行光信号放大。"
    )

    doc.add_heading("3.2 DWDM关键组件", level=2)
    doc.add_paragraph(
        "DWDM系统包含多种关键组件。"
        "其中OADM（光分插复用器）是DWDM网络中的核心节点设备，"
        "负责在不进行光电转换的情况下实现波长的上下路。"
        "EDFA（掺铒光纤放大器）属于光放大器的一种，"
        "是目前DWDM系统中使用最广泛的光信号放大设备。"
        "此外，DWDM系统依赖光纤进行信号传输，"
        "光纤属于传输介质，是整个光网络的物理基础。"
    )

    # ── 第4章 总结 ──
    doc.add_heading("第四章 技术总结", level=1)
    doc.add_paragraph(
        "综上所述，现代光传输网络以OTN为核心体系框架，"
        "融合了SDH的运维管理能力和DWDM的大容量传输能力。"
        "SDH提供STM-16（2.5G）和STM-64（10G）等标准速率接口，"
        "DWDM通过OADM实现灵活的波长调度，"
        "通过EDFA光放大器保证长距离传输质量，"
        "所有信号最终通过光纤这一传输介质进行承载。"
    )

    # 设置默认字体大小
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(12)

    doc.save(output_path)
    print(f"✅ 文档已生成: {output_path}")


if __name__ == "__main__":
    create_demo_document()
